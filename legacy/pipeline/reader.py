
from __future__ import annotations
from typing import Tuple, Dict
import docx2txt
from docx import Document

def read_docx_with_fallback(path: str, prefer: str = "python-docx", logger=None) -> Tuple[str, Dict]:
    meta = {"engine": None, "had_fallback": False, "tables_extracted": False}
    text = ""
    try:
        if prefer == "python-docx":
            doc = Document(path)
            parts = []
            # paragraphs
            parts.extend(p.text for p in doc.paragraphs)
            # tables
            for tbl in doc.tables:
                for row in tbl.rows:
                    parts.append(" | ".join(cell.text for cell in row.cells))
            text = "\n".join(parts)
            meta["engine"] = "python-docx"
            meta["tables_extracted"] = True
            if not text.strip():
                raise ValueError("empty text via python-docx")
        else:
            raise Exception("force fallback")
    except Exception:
    # fallback to docx2txt
        text = docx2txt.process(path) or ""
        meta["engine"] = "docx2txt"
        meta["had_fallback"] = True
    # hardening: if python-docx returned too little, try docx2txt and prefer the longer one
    if logger: logger.info(f"DOCX reader used: {meta['engine']} | fallback={meta['had_fallback']}")
    if meta["engine"] == "python-docx" and len(text) < 1000:
        alt = docx2txt.process(path) or ""
        if len(alt) > len(text):
            text = alt
            meta["engine"] = "docx2txt"
            meta["had_fallback"] = True
    if logger: logger.info(f"DOCX reader final: {meta['engine']} | fallback={meta['had_fallback']} | len={len(text)}")
    return text, meta

def repack_paragraphs(s: str) -> str:
    # single blank line between paragraphs
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    lines = [ln.strip() for ln in s.split("\n")]
    out, buf = [], []
    for ln in lines:
        if not ln:
            if buf:
                out.append(" ".join(buf).strip())
                buf = []
        else:
            buf.append(ln)
    if buf:
        out.append(" ".join(buf).strip())
    return "\n\n".join(out)
