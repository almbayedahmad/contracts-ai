from pathlib import Path
import re

p = Path(r".\app\ui_streamlit.py")
src = p.read_text(encoding="utf-8", errors="ignore")

# 1) طبع الأسطر وإزالة BOM
text = src.replace("\r\n", "\n").replace("\r", "\n")
if text.startswith("\ufeff"):
    text = text[1:]

# 2) كتلة الدوال النظيفة
readers_block = """
# --- Robust readers (clean) ---
def _read_text(file):
    \"\"\"Return textual content from an uploaded file-like object.\"\"\"
    suffix = Path(file.name).suffix.lower()
    data = file.read()
    if suffix == ".txt":
        try:
            return data.decode("utf-8")
        except Exception:
            return data.decode("latin-1", errors="ignore")
    elif suffix == ".docx":
        try:
            from docx import Document
            import io as _io
            doc = Document(_io.BytesIO(data))
            return "\\n".join(p.text for p in doc.paragraphs)
        except Exception:
            return ""
    # Fallback for any other extension
    try:
        return data.decode("utf-8")
    except Exception:
        return data.decode("latin-1", errors="ignore")


def _read_text_and_tables(file):
    \"\"\"Return (text, tables) from an uploaded file-like object.\"\"\"
    suffix = Path(file.name).suffix.lower()
    data = file.read()
    if suffix == ".docx":
        try:
            from docx import Document
            import io as _io
            doc = Document(_io.BytesIO(data))
            parts = []
            for p in doc.paragraphs:
                if p.text:
                    parts.append(p.text)
            tables = []
            for t in doc.tables:
                rows = []
                for r in t.rows:
                    rows.append([c.text.strip() for c in r.cells])
                tables.append(rows)
            return "\\n".join(parts), tables
        except Exception:
            return "", []
    else:
        try:
            text = data.decode("utf-8")
        except Exception:
            text = data.decode("latin-1", errors="ignore")
        return text, []
""".lstrip("\n")

# 3) احذف أي نسخ قديمة للدالتين ثم أدرج النسخة النظيفة مرة واحدة
pat_read_text = re.compile(r"(?ms)^\s*def\s+_read_text\([^\)]*\):.*?(?=^\s*def\s+|^\s*global\s+results_dir\b|$\Z)")
pat_read_tables = re.compile(r"(?ms)^\s*def\s+_read_text_and_tables\([^\)]*\):.*?(?=^\s*def\s+|^\s*global\s+results_dir\b|$\Z)")

text = pat_read_text.sub("", text)
text = pat_read_tables.sub("", text)

m = re.search(r"(?m)^\s*global\s+results_dir\b", text)
if m:
    insert_pos = m.start()
    text = text[:insert_pos].rstrip("\n") + "\n\n" + readers_block + "\n" + text[insert_pos:].lstrip("\n")
else:
    text = text.rstrip("\n") + "\n\n" + readers_block + "\n"

# 4) أزل أي بقايا backrefs حرفية
text = text.replace(r"\1", "")

# 5) مرّر سريعًا على except/finally اليتيمة وأدرج try/pass قبلها عند الحاجة
lines = text.split("\n")

def indent_of(s: str) -> int:
    return len(s) - len(s.lstrip(" "))

open_try = set()
fixed = []
for i, line in enumerate(lines):
    ind = indent_of(line)
    # إغلاق أي try مفتوحة أعمق من الإزاحة الحالية عند الرجوع بإزاحة
    for t in list(open_try):
        if t > ind:
            open_try.remove(t)

    stripped = line.lstrip()
    if re.match(r"^try:\s*$", stripped):
        open_try.add(ind)
    elif re.match(r"^(except\b|finally:)", stripped):
        if ind not in open_try:
            # إدراج try/pass صناعية قبل except
            fixed.append(" " * ind + "try:")
            fixed.append(" " * (ind + 4) + "pass")
        else:
            # اعتبرنا الـ try تمّت معالجتها عند هذه الإزاحة
            open_try.discard(ind)

    fixed.append(line)

text = "\n".join(fixed)
if not text.endswith("\n"):
    text += "\n"

p.write_text(text, encoding="utf-8")
print("Rewrote readers, normalized newlines, removed backrefs, and fixed orphan except/finally.")
