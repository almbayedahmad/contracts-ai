
from docx import Document
import fitz  # PyMuPDF

def read_docx_text(path: str) -> str:
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                parts.append(c.text)
    return "\n".join(parts)

def read_pdf_text(path: str) -> str:
    doc = fitz.open(path)
    texts = []
    for page in doc:
        texts.append(page.get_text("text"))
    return "\n".join(texts)

def read_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def read_docx_text_and_tables(path: str):
    from docx import Document
    doc = Document(path)
    parts = []
    tables = []
    # paragraphs
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    # tables
    for t in doc.tables:
        t_rows = []
        for r in t.rows:
            t_rows.append([c.text.strip() for c in r.cells])
            # also add to parts (for text search fallback)
            parts.append(" | ".join([c.text.strip() for c in r.cells if c.text]))
        tables.append(t_rows)
    return "\n".join(parts), tables
